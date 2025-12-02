import os
import tempfile
from pathlib import Path
from typing import List, BinaryIO, Dict, Any, Optional
import logging
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import uuid

# PDF parsing
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# DOCX parsing
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from vector_store import get_vector_store
from models import DocumentMetadata, IngestResponse

logger = logging.getLogger(__name__)


class DocumentIngester:
    """Handles document ingestion, parsing, chunking, and storage"""
    
    def __init__(self):
        self.vector_store = get_vector_store()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    async def ingest_file(
        self,
        file: BinaryIO,
        filename: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> IngestResponse:
        """
        Ingest a file into the vector store
        
        Args:
            file: File-like object
            filename: Original filename
            metadata: Optional metadata
            
        Returns:
            IngestResponse with status and IDs
        """
        try:
            # Extract text based on file type
            text = self._extract_text(file, filename)
            
            if not text.strip():
                raise ValueError("No text extracted from document")
            
            # Create document chunks
            documents = self._create_chunks(text, filename, metadata)
            
            # Add to vector store
            ids = self.vector_store.add_documents(documents)
            
            return IngestResponse(
                status="ok",
                ingested=len(ids),
                ids=ids,
                metadata={
                    "total_chunks": len(ids),
                    "source": filename,
                    "text_length": len(text)
                }
            )
            
        except Exception as e:
            logger.error(f"Error ingesting file {filename}: {e}")
            raise
    
    def _extract_text(self, file: BinaryIO, filename: str) -> str:
        """Extract text from various file formats"""
        ext = Path(filename).suffix.lower()
        
        if ext == ".pdf":
            return self._extract_pdf(file)
        elif ext in [".docx", ".doc"]:
            return self._extract_docx(file)
        elif ext == ".txt":
            return self._extract_text_file(file)
        else:
            # Try as text
            try:
                return self._extract_text_file(file)
            except:
                raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF"""
        if not HAS_PDF:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")
        
        try:
            pdf_reader = PdfReader(file)
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}\n")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def _extract_docx(self, file: BinaryIO) -> str:
        """Extract text from DOCX"""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Save to temp file (python-docx needs file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            
            try:
                doc = DocxDocument(tmp_path)
                text_parts = [paragraph.text for paragraph in doc.paragraphs]
                return "\n".join(text_parts)
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_text_file(self, file: BinaryIO) -> str:
        """Extract text from plain text file"""
        try:
            content = file.read()
            # Try UTF-8 first, fallback to latin-1
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                return content.decode("latin-1")
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            raise
    
    def _create_chunks(
        self,
        text: str,
        filename: str,
        metadata: Optional[DocumentMetadata]
    ) -> List[Document]:
        """Split text into chunks and create Document objects"""
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        base_metadata = {
            "source": filename,
            "doc_id": str(uuid.uuid4()),
        }
        
        if metadata:
            base_metadata.update({
                "source_name": metadata.source_name,
                "uploaded_by": metadata.uploaded_by,
                "timestamp": metadata.timestamp.isoformat(),
            })
            if metadata.geo:
                base_metadata["geo"] = metadata.geo
            if metadata.ward:
                base_metadata["ward"] = metadata.ward
            if metadata.tags:
                base_metadata["tags"] = ",".join(metadata.tags)
        
        for i, chunk in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata["chunk_id"] = i
            chunk_metadata["total_chunks"] = len(chunks)
            
            # Try to extract page number from chunk if present
            if "[Page " in chunk:
                try:
                    page_str = chunk.split("[Page ")[1].split("]")[0]
                    chunk_metadata["page"] = int(page_str)
                except:
                    pass
            
            documents.append(Document(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        return documents
    
    async def ingest_url(self, url: str, metadata: Optional[DocumentMetadata] = None) -> IngestResponse:
        """
        Ingest content from URL (future enhancement)
        
        TODO: Implement web scraping with beautiful soup or similar
        """
        raise NotImplementedError("URL ingestion not yet implemented")


# Singleton instance
_ingester_instance: Optional[DocumentIngester] = None


def get_ingester() -> DocumentIngester:
    """Get or create ingester singleton"""
    global _ingester_instance
    if _ingester_instance is None:
        _ingester_instance = DocumentIngester()
    return _ingester_instance
